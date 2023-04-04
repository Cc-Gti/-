# -*- coding: utf-8 -*-
"""
Created on Wed Nov  9 16:07:49 2022

@author: Lenovo
"""

from docx import Document
import pandas as pd
import numpy as np
import math
import matplotlib.pyplot as plt
import json
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from matplotlib.gridspec import GridSpec
# -*- coding: utf-8 -*-
import os
from scipy import stats
import time
from health_index import health_Score,health_Score_plt
from trend_func import Vis_temp_diag
from stat_func import  wear_nas_diag2
from log import log

#######替换NAN#############
def fillNaN_with_unifrand(df):
    """info:替换nan,用当前数组的均值和方差的随机数填充
    df:类型为Series、array"""
    a = df.values
    m = np.isnan(a) # mask of NaNs
    mu, sigma = df.mean(), df.std()
    a[m] = np.random.normal(mu, sigma, size=m.sum())
    return df 
#######替换0###############
def fillzero_with_unifrand(df):
    a = df.values
    m = np.where(a==0)[0] #返还的是tuple,所以添加[0]
    mu, sigma = df.mean(), df.std()
    a[m] = np.random.normal(mu, sigma, size=len(m))
    return df 
#######替换固定值val########
def fillval_with_unifrand(df,val):
    a = df.values
    m = np.where(a==val)[0] #返还的是tuple,所以添加[0]
    mu, sigma = df.mean(), df.std()
    a[m] = np.random.normal(mu, sigma, size=len(m)).reshape(-1,1)
    return df 

#######替换n倍标准差之外的所有数据######
def fill_nsigma_with_unifrand(df,n):
    a = df.values
   
    mu, sigma = df.mean(), df.std()
    m = np.where((a>mu+n*sigma)&(a<mu-n*sigma))[0] #返还的是tuple,所以添加[0]
    a[m] = np.random.normal(mu, sigma, size=len(m)).reshape(-1,1)
    return df 




def fill_most_num(series):####某一列替换重复次数最多的值
    
     mu, sigma = series.mean(), series.std()
     print( mu, sigma )
        
     max_count=series.value_counts()[0]#重复最多的数据的次数
     
     value=series.value_counts().index[0]#重复最多的数据
     
     repeat_index = series[series == value].index.tolist()
     
     print( series.loc[repeat_index].values.shape)
     
     a=series.values

     a[repeat_index]=np.random.normal(mu, sigma/3, size=len(repeat_index))

     return series    
 




    

from statsmodels.tsa.holtwinters import ExponentialSmoothing
plt.rcParams['font.sans-serif']=['SimHei']
plt.rcParams['axes.unicode_minus'] = False

def test_func1(df_data,arg):##
    """
    df:dataframe,
    arg:预测的参数，如水分、黏度等
    如果报错,
    
    
    """
    try:
        x_test=df_data[arg]
        kstest_result=stats.kstest(x_test.values, 'norm', (x_test.mean(), x_test.std()))
        if kstest_result[1]>0.01:
            ets = ExponentialSmoothing(x_test, trend='add', seasonal='add', seasonal_periods=2)
            r = ets.fit()
            pred = r.predict(start=len(x_test), end=len(x_test)+len(x_test)//3)
        
            pred_df=pd.DataFrame({
                'origin': x_test,
                #'fitted': r.fittedvalues,
                'pred': pred
            })
            pred_df['pred'][len(df_data[arg])-1]= pred_df['origin'][len(df_data[arg])-1]
            pred_df.plot(legend=True,figsize=(12,6))

            plt.show()
        elif math.isnan(kstest_result[1]):
            ets = ExponentialSmoothing(x_test, trend=None, seasonal='mul', seasonal_periods=2)
            r= ets.fit()
            pred = r.predict(start=len(x_test), end=len(x_test)+len(x_test)//3)
  
            pred_df=pd.DataFrame({
                 'origin': x_test,
                 #'fitted': r.fittedvalues,
                 'pred': pred
             })
            pred_df['pred'][len(df_data[arg])-1]= pred_df['origin'][len(df_data[arg])-1]        
            pred_df.plot(legend=True,figsize=(12,6))

            plt.show()        
        else: 
            ets = ExponentialSmoothing(x_test, trend=None, seasonal='add', seasonal_periods=2)
            r= ets.fit()
            pred = r.predict(start=len(x_test), end=len(x_test)+len(x_test)//3)
  
            pred_df=pd.DataFrame({
                 'origin': x_test,
                 #'fitted': r.fittedvalues,
                 'pred': pred
            })
            pred_df['pred'][len(df_data[arg])-1]= pred_df['origin'][len(df_data[arg])-1]        
            pred_df.plot(legend=True,figsize=(12,6))

            plt.show()
        print('test_func1 Done')     
    except Exception as e:
        log(e)
        print(e)
        pred_df=pd.DataFrame({
             'origin': df_data[arg],
             #'fitted': r.fittedvalues,
             'pred': x_test.iloc[0:len(x_test)//3]
        })
        print('test_func1 Fail')             
        # tag,validation_results,temp_pred_df = forcast_func(df_data,arg)   
        # pred_df=pd.DataFrame({
        #      'origin': df_data[arg],
        #      #'fitted': r.fittedvalues,
        #      'pred': temp_pred_df[arg]
        #  })
        # pred_df['pred'][len(df_data[arg])-1]= pred_df['origin'][len(df_data[arg])-1]               
    return pred_df







def datafile_to_word2(df,equip_name,vis,data_deal=True,method='linear',period='12H',facility_type='汽轮机',oil_type='rest',arg_x='黏度',arg_y='水活性',arg_z='油温'):
    """
    info:面向单台机组的总结报告，
    df:dataframe,数据源
    equip_name:监测机组的名称，如1#压缩机
    data_deal:数据处理开关,默认True,如果data_deal=True,则采用数据处理，默认方法是线性补充，是对空缺数据的补充。
    method:['linear','quadratic']
    period:['12H','6H','2H'],每多少小时取平均值
    vis:油品的黏度等级，用于判定黏度的范围,可查看configuration
    facility_type：['汽轮机','压缩机','液压系统','齿轮箱'],代表设备的类型,和磨损的指标有关,可查看stat_func.py
    oil_type：默认'rest'是水分的判定标准,可查看configuration
    arg_x:health_index的三维可视化属性之一
    arg_y:health_index的三维可视化属性之一
    arg_z:health_index的三维可视化属性之一
    
    return :分析报告docx,保存在项目路径下
    
    """
    document = Document()
    cur_path = os.path.split(os.path.realpath(__file__))[0]
    os.path.split(os.path.realpath(__file__))[0]
    word_tag=0  
    with open("./configuration.json", "r", encoding="utf-8") as f:
        content = json.load(f)      
    document.add_heading('%s油液分析报告'%equip_name, level=0)    
    document.add_heading('监测情况分析', level=1)
    
    #document.add_paragraph('监测情况分析','Title')
    try:
        df['时间'] = pd.to_datetime(df['时间'])
        df = df.sort_values(by = '时间')
        quarterly_data = df.resample(period, on='时间').mean()####.interpolate(method='quadratic')        
        if data_deal==True:
           quarterly_data =quarterly_data.interpolate(method=method)
        if data_deal==False:   
            pass
    except Exception:pass
    try:
        df['time'] = pd.to_datetime(df['time'])
        df = df.sort_values(by = 'time')
        quarterly_data = df.resample(period, on='时间').mean()####.interpolate(method='quadratic')
        if data_deal==True:
           quarterly_data =quarterly_data.interpolate(method=method)
        if data_deal==False:   
            pass        
      
    except Exception:pass  
    document.add_paragraph('    以下是%s机组%s至%s的监测数据情况,该系统使用的油品等级为%s,设备类型为%s,具体数据分析报告如下。'%(equip_name,quarterly_data.index[0],quarterly_data.index[-1],str(vis),facility_type))
    tag1,Vis_tag,Vis_comment,Temp_tag,Temp_tag,Vis_Temp_tag,Vis_Temp_comment=Vis_temp_diag(df,vis)
    document.add_heading('黏度', level=2)
    try:  
         
         word_tag=word_tag+1   

         
         describe_Seiries=quarterly_data['黏度'].describe().round(2)
         if Vis_tag=='黏度无明显变化趋势':         
             document.add_paragraph('    图%s分别是%s在线监测的黏度监测趋势图和分布情况。由图可知，当前油液黏度状态稳定，序列平稳无异常,部分异常值可视为传感器噪声。当前黏度均值为%s,标准差为%s,中位值为%s,最大值为%s,最小值为%s。\
                                   '%(word_tag,equip_name,describe_Seiries['mean'],describe_Seiries['std'],describe_Seiries['50%'],describe_Seiries['max'],describe_Seiries['min']))
         elif Vis_tag==content['Conclusion']['Vis_up'][0]: 
             document.add_paragraph('    图%s分别是%s在线监测的黏度监测趋势图和分布情况。由图可知，当前油液黏度有上升趋势，黏度升高将导致系统润滑不良，建议关注黏度变化情况。当前黏度均值为%s,标准差为%s,中位值为%s,最大值为%s,最小值为%s。\
                                   '%(word_tag,equip_name,describe_Seiries['mean'],describe_Seiries['std'],describe_Seiries['50%'],describe_Seiries['max'],describe_Seiries['min']))
         elif Vis_tag==content['Conclusion']['Vis_down'][0]:
             document.add_paragraph('    图%s分别是%s在线监测的黏度监测趋势图和分布情况。当前油液黏度有下降趋势，黏度下降将导致油膜强度下降，建议关注黏度变化情况。当前黏度均值为%s,标准差为%s,中位值为%s,最大值为%s,最小值为%s。\
                                   '%(word_tag,equip_name,describe_Seiries['mean'],describe_Seiries['std'],describe_Seiries['50%'],describe_Seiries['max'],describe_Seiries['min']))             
         elif Vis_tag==content['Conclusion']['Vis_high'][0]:
             document.add_paragraph('    图%s分别是%s在线监测的黏度监测趋势图和分布情况。当前油液黏度偏高，可能由于低温或氧化导致，建议关注黏度变化情况。当前黏度均值为%s,标准差为%s,中位值为%s,最大值为%s,最小值为%s。\
                                   '%(word_tag,equip_name,describe_Seiries['mean'],describe_Seiries['std'],describe_Seiries['50%'],describe_Seiries['max'],describe_Seiries['min']))             
         else:
             document.add_paragraph('    图%s分别是%s在线监测的黏度监测趋势图和分布情况。当前黏度均值为%s,标准差为%s,中位值为%s,最大值为%s,最小值为%s。\
                                   '%(word_tag,equip_name,describe_Seiries['mean'],describe_Seiries['std'],describe_Seiries['50%'],describe_Seiries['max'],describe_Seiries['min']))             
         fig = plt.figure(figsize=(12,5),
                          constrained_layout=True,#类似于tight_layout，使得各子图之间的距离自动调整【类似excel中行宽根据内容自适应】
                         )      
         gs = GridSpec(1, 5, figure=fig)                 
         ax1 = fig.add_subplot(gs[0, 0:3])
         #quarterly_data.loc[(quarterly_data['黏度']>0)]['黏度'].plot(title='%s'%equip_name)
         quarterly_data['黏度'].plot(title='%s'%equip_name)
         #plt.hlines(content['VIS'][str(vis)][0]+content['VIS'][str(vis)][1])##增加标准线
         plt.ylabel('黏度/cSt')

         ax2 = fig.add_subplot(gs[0, 3:5])
         quarterly_data.loc[(quarterly_data['黏度']>0)]['黏度'].plot(kind='hist',title='%s'%equip_name)
         plt.xlabel('黏度')
         plt.savefig(cur_path+'\png\黏度.png',bbox_inches ='tight',) 
         plt.show()
         inline_pic =document.add_picture(cur_path+'\png\黏度.png',width=Inches(6)) #,width=Inches(6)
         inline_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER#WD_PARAGRAPH_ALIGNMENT.CENTER
         # inline_shape.height = Cm(4)
         # inline_shape.width = Cm(8)
         paragraph =document.add_paragraph('图%s 黏度变化及分布情况'%word_tag) 
         paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER                    
    except Exception as e:
         word_tag=word_tag-1  
         log(e)
    
    try:  

         word_tag=word_tag+1   
         describe_Seiries=quarterly_data['40度黏度'].describe().round(2)
     
         if Vis_tag=='黏度无明显变化趋势':         
             document.add_paragraph('    图%s分别是%s在线监测的黏度监测趋势图和分布情况。由图可知，当前油液黏度状态稳定，序列平稳无异常,部分异常值可视为传感器噪声。当前黏度均值为%s,标准差为%s,中位值为%s,最大值为%s,最小值为%s。\
                                   '%(word_tag,equip_name,describe_Seiries['mean'],describe_Seiries['std'],describe_Seiries['50%'],describe_Seiries['max'],describe_Seiries['min']))
         elif Vis_tag==content['Conclusion']['Vis_up'][0]: 
             document.add_paragraph('    图%s分别是%s在线监测的黏度监测趋势图和分布情况。由图可知，当前油液黏度有上升趋势，黏度升高将导致系统润滑不良，建议关注黏度变化情况。当前黏度均值为%s,标准差为%s,中位值为%s,最大值为%s,最小值为%s。\
                                   '%(word_tag,equip_name,describe_Seiries['mean'],describe_Seiries['std'],describe_Seiries['50%'],describe_Seiries['max'],describe_Seiries['min']))
         elif Vis_tag==content['Conclusion']['Vis_down'][0]:
             document.add_paragraph('    图%s分别是%s在线监测的黏度监测趋势图和分布情况。当前油液黏度有下降趋势，黏度下降将导致油膜强度下降，建议关注黏度变化情况。当前黏度均值为%s,标准差为%s,中位值为%s,最大值为%s,最小值为%s。\
                                   '%(word_tag,equip_name,describe_Seiries['mean'],describe_Seiries['std'],describe_Seiries['50%'],describe_Seiries['max'],describe_Seiries['min']))             
         elif Vis_tag==content['Conclusion']['Vis_high'][0]:
             document.add_paragraph('    图%s分别是%s在线监测的黏度监测趋势图和分布情况。当前油液黏度偏高，可能由于低温或混氧化导致，建议关注黏度变化情况。当前黏度均值为%s,标准差为%s,中位值为%s,最大值为%s,最小值为%s。\
                                   '%(word_tag,equip_name,describe_Seiries['mean'],describe_Seiries['std'],describe_Seiries['50%'],describe_Seiries['max'],describe_Seiries['min']))             
         else:
             document.add_paragraph('    图%s分别是%s在线监测的黏度监测趋势图和分布情况。当前黏度均值为%s,标准差为%s,中位值为%s,最大值为%s,最小值为%s。\
                                   '%(word_tag,equip_name,describe_Seiries['mean'],describe_Seiries['std'],describe_Seiries['50%'],describe_Seiries['max'],describe_Seiries['min'])) 
         fig = plt.figure(figsize=(12,5),
                          constrained_layout=True,#类似于tight_layout，使得各子图之间的距离自动调整【类似excel中行宽根据内容自适应】
                         )      
         gs = GridSpec(1, 5, figure=fig)                 
         ax1 = fig.add_subplot(gs[0, 0:3])
         quarterly_data.loc[(quarterly_data['黏度']>0)]['黏度'].plot(title='%s'%equip_name)
         #plt.hlines(content['VIS'][str(vis)][0]+content['VIS'][str(vis)][1])##增加标准线
         plt.ylabel('黏度/cSt')

         ax2 = fig.add_subplot(gs[0, 3:5])
         quarterly_data.loc[(quarterly_data['黏度']>0)]['黏度'].plot(kind='hist',title='%s'%equip_name)
         plt.xlabel('黏度')
         plt.savefig(cur_path+'\png\黏度.png',bbox_inches ='tight',) 
         plt.show()
         inline_pic =document.add_picture(cur_path+'\png\黏度.png',width=Inches(6)) #,width=Inches(6)
         inline_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER#WD_PARAGRAPH_ALIGNMENT.CENTER
         # inline_shape.height = Cm(4)
         # inline_shape.width = Cm(8)
         paragraph =document.add_paragraph('图%s 黏度变化及分布情况'%word_tag) 
         paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER         
    except Exception as e:
         word_tag=word_tag-1   
         log(e)
    document.add_heading('温度', level=2)    
    try:  

         word_tag=word_tag+1
         describe_Seiries=quarterly_data['温度'].describe().round(2)
         if Temp_tag=='温度无明显变化趋势':         
             document.add_paragraph('    图%s分别是%s在线监测的温度监测趋势图和分布情况。由图可知，当前温度状态稳定，序列平稳无异常,部分异常值可视为传感器噪声。当前温度均值为%s,标准差为%s,中位值为%s,最大值为%s,最小值为%s。\
                                   '%(word_tag,equip_name,describe_Seiries['mean'],describe_Seiries['std'],describe_Seiries['50%'],describe_Seiries['max'],describe_Seiries['min']))
         elif Vis_tag==content['Conclusion']['Temp_up'][0]: 
             document.add_paragraph('    图%s分别是%s在线监测的温度监测趋势图和分布情况。由图可知，当前油液温度有上升趋势，温度升高将导致系统润滑不良，建议关注温度变化情况。当前均值为%s,标准差为%s,中位值为%s,最大值为%s,最小值为%s。\
                                   '%(word_tag,equip_name,describe_Seiries['mean'],describe_Seiries['std'],describe_Seiries['50%'],describe_Seiries['max'],describe_Seiries['min']))                                   
         else:
             document.add_paragraph('    图%s分别是%s在线监测的黏度监测趋势图和分布情况。当前温度均值为%s,标准差为%s,中位值为%s,最大值为%s,最小值为%s。\
                                   '%(word_tag,equip_name,describe_Seiries['mean'],describe_Seiries['std'],describe_Seiries['50%'],describe_Seiries['max'],describe_Seiries['min']))          
         fig = plt.figure(figsize=(12,5),
                          constrained_layout=True,#类似于tight_layout，使得各子图之间的距离自动调整【类似excel中行宽根据内容自适应】
                         )      
         gs = GridSpec(1, 5, figure=fig)                 
         ax1 = fig.add_subplot(gs[0, 0:3])
         quarterly_data.loc[(quarterly_data['温度']>0)]['温度'].plot(title='%s'%equip_name)
         plt.ylabel('温度/℃')

         ax2 = fig.add_subplot(gs[0, 3:5])
         quarterly_data.loc[(quarterly_data['温度']>0)]['温度'].plot(kind='hist',title='%s'%equip_name)
         plt.xlabel('温度')
         plt.savefig(cur_path+'\png\温度.png',bbox_inches ='tight',) 
         plt.show()
         inline_pic =document.add_picture(cur_path+'\png\温度.png',width=Inches(6)) #,width=Inches(6)
         inline_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER#WD_PARAGRAPH_ALIGNMENT.CENTER
         # inline_shape.height = Cm(4)
         # inline_shape.width = Cm(8)
         paragraph =document.add_paragraph('图%s 温度变化及分布情况'%word_tag) 
         paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER       
    except Exception as e:
         word_tag=word_tag-1   
         log(e)
    try:  
         word_tag=word_tag+1   
         describe_Seiries=quarterly_data['油温'].describe().round(2)
         if Temp_tag=='温度无明显变化趋势':         
             document.add_paragraph('    图%s分别是%s在线监测的温度监测趋势图和分布情况。由图可知，当前温度状态稳定，序列平稳无异常,部分异常值可视为传感器噪声。当前温度均值为%s,标准差为%s,中位值为%s,最大值为%s,最小值为%s。\
                                   '%(word_tag,equip_name,describe_Seiries['mean'],describe_Seiries['std'],describe_Seiries['50%'],describe_Seiries['max'],describe_Seiries['min']))
         elif Vis_tag==content['Conclusion']['Temp_up'][0]: 
             document.add_paragraph('    图%s分别是%s在线监测的温度监测趋势图和分布情况。由图可知，当前油液温度有上升趋势，表明系统可能存在冷却水流量不足或摩擦副异常磨损；建议关注设备运行情况，建议关注温度变化情况。当前均值为%s,标准差为%s,中位值为%s,最大值为%s,最小值为%s。\
                                   '%(word_tag,equip_name,describe_Seiries['mean'],describe_Seiries['std'],describe_Seiries['50%'],describe_Seiries['max'],describe_Seiries['min']))
         else:
             document.add_paragraph('    图%s分别是%s在线监测的温度监测趋势图和分布情况。当前温度均值为%s,标准差为%s,中位值为%s,最大值为%s,最小值为%s。\
                                   '%(word_tag,equip_name,describe_Seiries['mean'],describe_Seiries['std'],describe_Seiries['50%'],describe_Seiries['max'],describe_Seiries['min']))     
         fig = plt.figure(figsize=(12,5),
                          constrained_layout=True,#类似于tight_layout，使得各子图之间的距离自动调整【类似excel中行宽根据内容自适应】
                         )      
         gs = GridSpec(1, 5, figure=fig)                 
         ax1 = fig.add_subplot(gs[0, 0:3])
         quarterly_data.loc[(quarterly_data['油温']>0)]['油温'].plot(title='%s'%equip_name)
         #plt.hlines(content['VIS'][str(vis)][0]+content['VIS'][str(vis)][1])##增加标准线
         plt.ylabel('油温/℃')

         ax2 = fig.add_subplot(gs[0, 3:5])
         quarterly_data.loc[(quarterly_data['油温']>0)]['油温'].plot(kind='hist',title='%s'%equip_name)
         plt.xlabel('油温')
         plt.savefig(cur_path+'\png\油温.png',bbox_inches ='tight',) 
         plt.show()
         inline_pic =document.add_picture(cur_path+'\png\油温.png',width=Inches(6)) #,width=Inches(6)
         inline_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER#WD_PARAGRAPH_ALIGNMENT.CENTER
         # inline_shape.height = Cm(4)
         # inline_shape.width = Cm(8)
         paragraph =document.add_paragraph('图%s 油温变化及分布情况'%word_tag)
         paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER              
    except Exception as e:
         word_tag=word_tag-1   
         log(e)
    document.add_heading('水活性', level=2)       
    try:  
         word_tag=word_tag+1
         describe_Seiries=quarterly_data['水活性'].describe().round(2)  
         #####缺少判定条件####
         #####缺少判定条件####
         
         document.add_paragraph('    图%s是%s测得水活性(水分)监测趋势图。由图可知，其水活性低于报警值，说明机组当前状态下在用油溶解水处于低位，未发生外来水分入侵现象。当前水活性均值为%s,标准差为%s,中位值位%s,最大值为%s,最小值为%s。'\
                               %(word_tag,equip_name,describe_Seiries['mean'],describe_Seiries['std'],describe_Seiries['50%'],describe_Seiries['max'],describe_Seiries['min']))        
         fig = plt.figure(figsize=(12,5),
                          constrained_layout=True,#类似于tight_layout，使得各子图之间的距离自动调整【类似excel中行宽根据内容自适应】
                         )      
         gs = GridSpec(1, 5, figure=fig)                 
         ax1 = fig.add_subplot(gs[0, 0:3])
         quarterly_data.loc[(quarterly_data['水活性']>0)]['水活性'].plot(title='%s'%equip_name)
         #plt.hlines(content['VIS'][str(vis)][0]+content['VIS'][str(vis)][1])##增加标准线
         plt.ylabel('水活性/aw')

         ax2 = fig.add_subplot(gs[0, 3:5])
         quarterly_data.loc[(quarterly_data['水活性']>0)]['水活性'].plot(kind='hist',title='%s'%equip_name)
         plt.xlabel('水活性')
         plt.savefig(cur_path+'\png\水活性.png',bbox_inches ='tight',) 
         plt.show()
         inline_pic =document.add_picture(cur_path+'\png\水活性.png',width=Inches(6)) #,width=Inches(6)
         inline_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER#WD_PARAGRAPH_ALIGNMENT.CENTER         
         paragraph =document.add_paragraph('图%s 水活性序列变化情况'%word_tag)
         paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER            

    except Exception as e:
         word_tag=word_tag-1   
         log(e) 
    
    tag,NAS_tag,NAS_comment,Fe_tag, Fe_comment ,nFe_tag,nFe_comment =wear_nas_diag2(df,facility_type)
    
    document.add_heading('铁磁颗粒', level=2)    
    
    try:  
  
         word_tag=word_tag+1
         describe_Seiries=quarterly_data[['总铁磁颗粒','铁磁细颗粒','铁磁粗颗粒']].describe().round(2)      
         
         if Fe_tag=='近期铁磁颗粒磨损情况正常':
         
             document.add_paragraph('    图%s是%s测得铁磁颗粒的监测趋势及分布图。由图可知，各尺寸铁磁颗粒低于报警值，说明机组未产生磨损。以下是各个尺寸颗粒的监测统计情况。'%(word_tag,equip_name))        
        
         elif  Fe_tag==content['Conclusion']['Fe_up'][0]:
             
             document.add_paragraph('    图%s是%s测得铁磁颗粒的监测趋势及分布图。由图可知，系统监测到铁磁部件异常磨损颗粒。以下是各个尺寸颗粒的监测统计情况。'%(word_tag,equip_name)) 
             
         else:
             
             document.add_paragraph('    图%s是%s测得铁磁颗粒的监测趋势及分布图，以下是各个尺寸颗粒的监测统计情况。'%(word_tag,equip_name))  
             
         document.add_paragraph('  当前总铁磁颗粒均值为%s,标准差为%s,中位值位%s,最大值为%s,最小值为%s;'\
                                %(describe_Seiries['总铁磁颗粒']['mean'],describe_Seiries['总铁磁颗粒']['std'],describe_Seiries['总铁磁颗粒']['50%'],describe_Seiries['总铁磁颗粒']['max'],describe_Seiries['总铁磁颗粒']['min']))
         
         document.add_paragraph('  当前铁磁细颗粒均值为%s,标准差为%s,中位值位%s,最大值为%s,最小值为%s;'\
                                %(describe_Seiries['铁磁细颗粒']['mean'],describe_Seiries['铁磁细颗粒']['std'],describe_Seiries['铁磁细颗粒']['50%'],describe_Seiries['铁磁细颗粒']['max'],describe_Seiries['铁磁细颗粒']['min']))  
         document.add_paragraph('  当前铁磁粗颗粒均值为%s,标准差为%s,中位值位%s,最大值为%s,最小值为%s。'\
                                %(describe_Seiries['铁磁粗颗粒']['mean'],describe_Seiries['铁磁粗颗粒']['std'],describe_Seiries['铁磁粗颗粒']['50%'],describe_Seiries['铁磁粗颗粒']['max'],describe_Seiries['铁磁粗颗粒']['min']))           
             
         fig = plt.figure(figsize=(12,6),
                          constrained_layout=True,#类似于tight_layout，使得各子图之间的距离自动调整【类似excel中行宽根据内容自适应】
                         )      
         gs = GridSpec(3, 5, figure=fig)                 
         ax1 = fig.add_subplot(gs[0, 0:3])
         quarterly_data['总铁磁颗粒'].plot(title='%s'%equip_name)
         #plt.hlines(content['VIS'][str(vis)][0]+content['VIS'][str(vis)][1])##增加标准线
         plt.ylabel('总铁磁颗粒')

         ax2 = fig.add_subplot(gs[0, 3:5])
         quarterly_data['总铁磁颗粒'].plot(kind='hist',title='%s'%equip_name)
         plt.xlabel('总铁磁颗粒')

         ax3 = fig.add_subplot(gs[1, 0:3])
         quarterly_data['铁磁细颗粒'].plot()
         #plt.hlines(content['VIS'][str(vis)][0]+content['VIS'][str(vis)][1])##增加标准线
         plt.ylabel('铁磁细颗粒')  
         
         ax4 = fig.add_subplot(gs[1, 3:5])
         quarterly_data['铁磁细颗粒'].plot(kind='hist')
         plt.xlabel('铁磁细颗粒')  
         
         ax5 = fig.add_subplot(gs[2, 0:3])
         quarterly_data['铁磁粗颗粒'].plot()
         plt.ylabel('铁磁粗颗粒')     
         
         ax6 = fig.add_subplot(gs[2, 3:5])
         quarterly_data['铁磁粗颗粒'].plot(kind='hist')
         plt.xlabel('铁磁粗颗粒')           

         plt.savefig(cur_path+'\png\铁磁颗粒.png',bbox_inches ='tight',) 
         inline_pic =document.add_picture(cur_path+'\png\铁磁颗粒.png',width=Inches(6)) #,width=Inches(6)
         inline_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER#WD_PARAGRAPH_ALIGNMENT.CENTER         
         paragraph =document.add_paragraph('图%s 铁磁颗粒趋势及分布情况'%word_tag) 
         paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER            
    except Exception as e:
         word_tag=word_tag-1   
         log(e)    
    try:  
         
         word_tag=word_tag+1
         describe_Seiries=quarterly_data[['Fe(70~100um)','Fe(100~150um)','Fe(>150um)']].describe().round(2)  
         
         if Fe_tag=='近期铁磁颗粒磨损情况正常':
         
             document.add_paragraph('    图%s是%s测得铁磁颗粒的监测趋势及分布图。由图可知，各尺寸铁磁颗粒低于报警值，说明机组未产生磨损。以下是各个尺寸颗粒的监测统计情况。'%(word_tag,equip_name))        
        
         elif  Fe_tag==content['Conclusion']['Fe_up'][0]:
             
             document.add_paragraph('    图%s是%s测得铁磁颗粒的监测趋势及分布图。由图可知，系统监测到铁磁部件异常磨损颗粒。以下是各个尺寸颗粒的监测统计情况。'%(word_tag,equip_name)) 
             
         else:
             
             document.add_paragraph('    图%s是%s测得铁磁颗粒的监测趋势及分布图，以下是各个尺寸颗粒的监测统计情况。'%(word_tag,equip_name))  
       
         document.add_paragraph('  当前铁磁颗粒Fe(70~100um)均值为%s,标准差为%s,中位值位%s,最大值为%s,最小值为%s;'\
                                %(describe_Seiries['Fe(70~100um)']['mean'],describe_Seiries['Fe(70~100um)']['std'],describe_Seiries['Fe(70~100um)']['50%'],describe_Seiries['Fe(70~100um)']['max'],describe_Seiries['Fe(70~100um)']['min']))
         
         document.add_paragraph('  当前铁磁颗粒Fe(100~150um)均值为%s,标准差为%s,中位值位%s,最大值为%s,最小值为%s;'\
                                %(describe_Seiries['Fe(100~150um)']['mean'],describe_Seiries['Fe(100~150um)']['std'],describe_Seiries['Fe(100~150um)']['50%'],describe_Seiries['Fe(100~150um)']['max'],describe_Seiries['Fe(100~150um)']['min']))  
         document.add_paragraph('  当前铁磁颗粒Fe(>150um)均值为%s,标准差为%s,中位值位%s,最大值为%s,最小值为%s。'\
                                %(describe_Seiries['Fe(>150um)']['mean'],describe_Seiries['Fe(>150um)']['std'],describe_Seiries['Fe(>150um)']['50%'],describe_Seiries['Fe(>150um)']['max'],describe_Seiries['Fe(>150um)']['min']))              
         fig = plt.figure(figsize=(12,6),
                          constrained_layout=True,#类似于tight_layout，使得各子图之间的距离自动调整【类似excel中行宽根据内容自适应】
                         )      
         gs = GridSpec(3, 5, figure=fig)                 
         ax1 = fig.add_subplot(gs[0, 0:3])
         quarterly_data['Fe(>150um)'].plot(title='%s'%equip_name)
         #plt.hlines(content['VIS'][str(vis)][0]+content['VIS'][str(vis)][1])##增加标准线
         plt.ylabel('Fe(>150um)')

         ax2 = fig.add_subplot(gs[0, 3:5])
         quarterly_data['Fe(>150um)'].plot(kind='hist',title='%s'%equip_name)
         plt.xlabel('Fe(>150um)')

         ax3 = fig.add_subplot(gs[1, 0:3])
         quarterly_data['Fe(100~150um)'].plot()
         #plt.hlines(content['VIS'][str(vis)][0]+content['VIS'][str(vis)][1])##增加标准线
         plt.ylabel('Fe(100~150um)')  
         
         ax4 = fig.add_subplot(gs[1, 3:5])
         quarterly_data['Fe(100~150um)'].plot(kind='hist')
         plt.xlabel('Fe(100~150um)')  
         
         ax5 = fig.add_subplot(gs[2, 0:3])
         quarterly_data['Fe(70~100um)'].plot()
         plt.ylabel('Fe(100~150um)')     
         
         ax6 = fig.add_subplot(gs[2, 3:5])
         quarterly_data['Fe(70~100um)'].plot(kind='hist')
         plt.xlabel('Fe(70~100um)')           

         plt.savefig(cur_path+'\png\铁磁颗粒.png',bbox_inches ='tight',) 
         inline_pic =document.add_picture(cur_path+'\png\铁磁颗粒.png',width=Inches(6)) #,width=Inches(6)
         inline_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER#WD_PARAGRAPH_ALIGNMENT.CENTER         
         paragraph =document.add_paragraph('图%s 铁磁颗粒趋势及分布情况'%word_tag) 
         paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
         word_tag=word_tag-1   
         log(e)    
    document.add_heading('非铁磁颗粒', level=2)    
    try:  
         word_tag=word_tag+1
         describe_Seiries=quarterly_data[['非铁磁颗粒(200um~300um)','非铁磁颗粒(300um~400um)','非铁磁颗粒(400um~)']].describe().round(2)    
         if nFe_tag=='近期非铁磁颗粒磨损情况正常':
         
             document.add_paragraph('    图%s是%s测得非铁磁颗粒的监测趋势及分布图。由图可知，各尺寸非铁磁颗粒低于报警值，说明机组未产生磨损。以下是各个尺寸颗粒的监测统计情况。'%(word_tag,equip_name))        
        
         elif  nFe_tag==content['Conclusion']['nFe_up'][0]:
             
             document.add_paragraph('    图%s是%s测得非铁磁颗粒的监测趋势及分布图。由图可知，系统监测到铁磁部件异常磨损颗粒。以下是各个尺寸颗粒的监测统计情况。'%(word_tag,equip_name)) 
             
         else:
             
             document.add_paragraph('    图%s是%s测得非铁磁颗粒的监测趋势及分布图，以下是各个尺寸颗粒的监测统计情况。'%(word_tag,equip_name))           
         
         document.add_paragraph('  当前非铁磁颗粒(200um~300um)均值为%s,标准差为%s,中位值位%s,最大值为%s,最小值为%s;'\
                                %(describe_Seiries['非铁磁颗粒(200um~300um)']['mean'],describe_Seiries['非铁磁颗粒(200um~300um)']['std'],describe_Seiries['非铁磁颗粒(200um~300um)']['50%'],describe_Seiries['非铁磁颗粒(200um~300um)']['max'],describe_Seiries['非铁磁颗粒(200um~300um)']['min']))
         
         document.add_paragraph('  当前非铁磁颗粒(300um~400um)均值为%s,标准差为%s,中位值位%s,最大值为%s,最小值为%s;'\
                                %(describe_Seiries['非铁磁颗粒(300um~400um)']['mean'],describe_Seiries['非铁磁颗粒(300um~400um)']['std'],describe_Seiries['非铁磁颗粒(300um~400um)']['50%'],describe_Seiries['非铁磁颗粒(300um~400um)']['max'],describe_Seiries['非铁磁颗粒(300um~400um)']['min']))  
         document.add_paragraph('  当前非铁磁颗粒(400um~)均值为%s,标准差为%s,中位值位%s,最大值为%s,最小值为%s。'\
                                %(describe_Seiries['非铁磁颗粒(400um~)']['mean'],describe_Seiries['非铁磁颗粒(400um~)']['std'],describe_Seiries['非铁磁颗粒(400um~)']['50%'],describe_Seiries['非铁磁颗粒(400um~)']['max'],describe_Seiries['非铁磁颗粒(400um~)']['min']))                
         fig = plt.figure(figsize=(12,6),
                          constrained_layout=True,#类似于tight_layout，使得各子图之间的距离自动调整【类似excel中行宽根据内容自适应】
                         )      
         gs = GridSpec(3, 5, figure=fig)                 
         ax1 = fig.add_subplot(gs[0, 0:3])
         quarterly_data['非铁磁颗粒(200um~300um)'].plot(title='%s'%equip_name)
         #plt.hlines(content['VIS'][str(vis)][0]+content['VIS'][str(vis)][1])##增加标准线
         plt.ylabel('非铁磁颗粒(200um~300um)')

         ax2 = fig.add_subplot(gs[0, 3:5])
         quarterly_data['非铁磁颗粒(200um~300um)'].plot(kind='hist',title='%s'%equip_name)
         plt.xlabel('非铁磁颗粒(200um~300um)')

         ax3 = fig.add_subplot(gs[1, 0:3])
         quarterly_data['非铁磁颗粒(300um~400um)'].plot()
         #plt.hlines(content['VIS'][str(vis)][0]+content['VIS'][str(vis)][1])##增加标准线
         plt.ylabel('非铁磁颗粒(300um~400um)')  
         
         ax4 = fig.add_subplot(gs[1, 3:5])
         quarterly_data['非铁磁颗粒(300um~400um)'].plot(kind='hist')
         plt.xlabel('非铁磁颗粒(300um~400um)')  
         
         ax5 = fig.add_subplot(gs[2, 0:3])
         quarterly_data['非铁磁颗粒(400um~)'].plot()
         plt.ylabel('非铁磁颗粒(400um~)')     
         
         ax6 = fig.add_subplot(gs[2, 3:5])
         quarterly_data['非铁磁颗粒(400um~)'].plot(kind='hist')
         plt.xlabel('非铁磁颗粒(400um~)')           

         plt.savefig(cur_path+'\png\铁磁颗粒.png',bbox_inches ='tight',) 
         inline_pic =document.add_picture(cur_path+'\png\铁磁颗粒.png',width=Inches(6)) #,width=Inches(6)
         inline_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER#WD_PARAGRAPH_ALIGNMENT.CENTER         
         paragraph =document.add_paragraph('图%s 非铁磁颗粒趋势及分布情况'%word_tag) 
         paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER            
    except Exception as e:
         word_tag=word_tag-1   
         log(e) 

    try:  
         word_tag=word_tag+1
         describe_Seiries=quarterly_data[['nFe(200~300um)','nFe(300~400um)','nFe(>400um)']].describe().round(2)      
         #describe_Seiries=quarterly_data[['非铁磁颗粒(200um~300um)','非铁磁颗粒(300um~400um)','非铁磁颗粒(400um~)']].describe().round(2)      
         if nFe_tag=='近期非铁磁颗粒磨损情况正常':
         
             document.add_paragraph('    图%s是%s测得非铁磁颗粒的监测趋势及分布图。由图可知，各尺寸非铁磁颗粒低于报警值，说明机组未产生磨损。以下是各个尺寸颗粒的监测统计情况。'%(word_tag,equip_name))        
        
         elif  nFe_tag==content['Conclusion']['nFe_up'][0]:
             
             document.add_paragraph('    图%s是%s测得非铁磁颗粒的监测趋势及分布图。由图可知，系统监测到非铁磁部件异常磨损颗粒。以下是各个尺寸颗粒的监测统计情况。'%(word_tag,equip_name)) 
             
         else:
             
             document.add_paragraph('    图%s是%s测得非铁磁颗粒的监测趋势及分布图，以下是各个尺寸颗粒的监测统计情况。'%(word_tag,equip_name))  
             
         document.add_paragraph('  当前非铁磁颗粒nFe(200~300um)均值为%s,标准差为%s,中位值位%s,最大值为%s,最小值为%s;'\
                                %(describe_Seiries['nFe(200~300um)']['mean'],describe_Seiries['nFe(200~300um)']['std'],describe_Seiries['nFe(200~300um)']['50%'],describe_Seiries['nFe(200~300um)']['max'],describe_Seiries['nFe(200~300um)']['min']))
         
         document.add_paragraph('  当前非铁磁颗粒nFe(300~400um)均值为%s,标准差为%s,中位值位%s,最大值为%s,最小值为%s;'\
                                %(describe_Seiries['nFe(300~400um)']['mean'],describe_Seiries['nFe(300~400um)']['std'],describe_Seiries['nFe(300~400um)']['50%'],describe_Seiries['nFe(300~400um)']['max'],describe_Seiries['nFe(300~400um)']['min']))  
         document.add_paragraph('  当前非铁磁颗粒nFe(>400um)均值为%s,标准差为%s,中位值位%s,最大值为%s,最小值为%s。'\
                                %(describe_Seiries['nFe(>400um)']['mean'],describe_Seiries['nFe(>400um)']['std'],describe_Seiries['nFe(>400um)']['50%'],describe_Seiries['nFe(>400um)']['max'],describe_Seiries['nFe(>400um)']['min']))              
         fig = plt.figure(figsize=(12,6),
                          constrained_layout=True,#类似于tight_layout，使得各子图之间的距离自动调整【类似excel中行宽根据内容自适应】
                         )      
         gs = GridSpec(3, 5, figure=fig)                 
         ax1 = fig.add_subplot(gs[0, 0:3])
         quarterly_data['nFe(200~300um)'].plot(title='%s'%equip_name)
         #plt.hlines(content['VIS'][str(vis)][0]+content['VIS'][str(vis)][1])##增加标准线
         plt.ylabel('nFe(200~300um)')

         ax2 = fig.add_subplot(gs[0, 3:5])
         quarterly_data['nFe(300~400um)'].plot(kind='hist',title='%s'%equip_name)
         plt.xlabel('nFe(300~400um)')

         ax3 = fig.add_subplot(gs[1, 0:3])
         quarterly_data['nFe(300~400um)'].plot()
         #plt.hlines(content['VIS'][str(vis)][0]+content['VIS'][str(vis)][1])##增加标准线
         plt.ylabel('nFe(300~400um)')  
         
         ax4 = fig.add_subplot(gs[1, 3:5])
         quarterly_data['nFe(300~400um)'].plot(kind='hist')
         plt.xlabel('nFe(300~400um)')  
         
         ax5 = fig.add_subplot(gs[2, 0:3])
         quarterly_data['nFe(>400um)'].plot()
         plt.ylabel('nFe(>400um)')     
         
         ax6 = fig.add_subplot(gs[2, 3:5])
         quarterly_data['nFe(>400um)'].plot(kind='hist')
         plt.xlabel('nFe(>400um)')           

         plt.savefig(cur_path+'\png\非铁磁颗粒.png',bbox_inches ='tight',) 
         inline_pic =document.add_picture(cur_path+'\png\非铁磁颗粒.png',width=Inches(6)) #,width=Inches(6)
         inline_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER#WD_PARAGRAPH_ALIGNMENT.CENTER         
         paragraph =document.add_paragraph('图%s 非铁磁颗粒趋势及分布情况'%word_tag) 
         paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
         word_tag=word_tag-1   
         log(e)   
    document.add_heading('污染度', level=2)
    try:  
         word_tag=word_tag+1
         describe_Seiries=quarterly_data['NAS等级'].describe().round(2)

         if NAS_tag=='近期污染情况正常':
             
             document.add_paragraph('图%s分别是%s的在用污染度趋势图与分布图。由图可知，系统污染状态正常，当前污染度均值为%s,标准差为%s,中位值位%s,最大值为%s,最小值为%s。'\
                                   %(word_tag,equip_name,describe_Seiries['mean'].round(0),describe_Seiries['std'],describe_Seiries['50%'].round(0),describe_Seiries['max'].round(0),describe_Seiries['min'].round(0))) 
         elif NAS_tag==content['Conclusion']['NAS_up'][0]:
            
             document.add_paragraph('图%s分别是%s的在用污染度趋势图与分布图。由图可知，污染度有上升趋势，当前污染度均值为%s,标准差为%s,中位值位%s,最大值为%s,最小值为%s。'\
                                   %(word_tag,equip_name,describe_Seiries['mean'].round(0),describe_Seiries['std'],describe_Seiries['50%'].round(0),describe_Seiries['max'].round(0),describe_Seiries['min'].round(0))) 
         else:    
              document.add_paragraph('图%s分别是%s的在用污染度趋势图与分布图,当前污染度均值为%s,标准差为%s,中位值位%s,最大值为%s,最小值为%s。'\
                                   %(word_tag,equip_name,describe_Seiries['mean'].round(0),describe_Seiries['std'],describe_Seiries['50%'].round(0),describe_Seiries['max'].round(0),describe_Seiries['min'].round(0))) 
         fig = plt.figure(figsize=(12,6),
                          constrained_layout=True,#类似于tight_layout，使得各子图之间的距离自动调整【类似excel中行宽根据内容自适应】
                         )      
         gs = GridSpec(1, 5, figure=fig)                 
         ax1 = fig.add_subplot(gs[0, 0:3])
         quarterly_data['NAS等级'].plot(title='%s'%equip_name)#'4umISO','6umISO','14umISO','21umISO',
         #plt.hlines(content['VIS'][str(vis)][0]+content['VIS'][str(vis)][1])##增加标准线
         plt.ylabel('NAS等级')

         ax2 = fig.add_subplot(gs[0, 3:5])
         quarterly_data['NAS等级'].plot(kind='hist',title='%s'%equip_name)#'4umISO','6umISO','14umISO','21umISO',
         plt.xlabel('NAS等级')
         plt.savefig(cur_path+'\\png\\NAS等级.png',bbox_inches ='tight') 
         plt.show()
         inline_pic =document.add_picture(cur_path+'\\png\\NAS等级.png',width=Inches(6)) #,width=Inches(6)
         inline_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER#WD_PARAGRAPH_ALIGNMENT.CENTER         
         paragraph =document.add_paragraph('图%s 污染度趋势及分布情况'%word_tag) 
         paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER            
    except Exception as e:
         word_tag=word_tag-1   
         log(e)   
    try:     
        word_tag=word_tag+1
        document.add_heading('预测分析', level=1)
        quarterly_data['时间']=quarterly_data.index
    
        if '黏度' in quarterly_data.columns:
            Vis_pred_df=test_func1(quarterly_data,'黏度')
            Vis_pred_df.plot(figsize=(12,5),title='%s'%equip_name) 
            plt.ylabel('黏度/cSt')     
            plt.savefig(cur_path+'//png//forcast_func_黏度.png',bbox_inches ='tight')              
        if '40度黏度' in  quarterly_data.columns:
            Vis_pred_df=test_func1(quarterly_data,'40度黏度')    
            Vis_pred_df.plot(figsize=(12,5),title='%s'%equip_name) 
            plt.ylabel('40度黏度/cSt')     
            plt.savefig(cur_path+'//png//forcast_func_黏度.png',bbox_inches ='tight')                
        if '水活性' in quarterly_data.columns:
            Water_pred_df=test_func1(quarterly_data,'水活性')   
            Water_pred_df.plot(figsize=(12,5),title='%s'%equip_name) 
            plt.ylabel('水活性/aw')        
            plt.savefig(cur_path+'//png//forcast_func_水活性.png',bbox_inches ='tight') 
        if '温度' in quarterly_data.columns:             
            Temp_pred_df=test_func1(quarterly_data,'温度')   
            Temp_pred_df.plot(figsize=(12,5),title='%s'%equip_name) 
            plt.ylabel('温度/℃')
            plt.savefig(cur_path+'//png//forcast_func_温度.png',bbox_inches ='tight') 
        if  '油温' in quarterly_data.columns:      
            Temp_pred_df=test_func1(quarterly_data,'油温')   
            Temp_pred_df.plot(figsize=(12,5),title='%s'%equip_name) 
            plt.ylabel('油温/℃')
            plt.savefig(cur_path+'//png//forcast_func_油温.png',bbox_inches ='tight')     
            
        document.add_paragraph('图%s是%s的在用油黏度、温度及水活性的趋势预测情况。由图可知，'%(word_tag,equip_name))
    
        if (content['VIS'][str(vis)][0]-content['VIS'][str(vis)][1]>Vis_pred_df['pred'].dropna()).all() or (Vis_pred_df['pred'].dropna()>content['VIS'][str(vis)][0]+content['VIS'][str(vis)][1]).all():
            
            document.add_paragraph('预测黏度未来有超过阈值的趋势;建议加强关注，必要时检查系统用油情况。')
        else: document.add_paragraph('预测黏度未来仍处于正常范围;')   
        if (Water_pred_df['pred'].dropna()>content['AW']['rest'][0]).all(): 
            document.add_paragraph('预测水分未来有超过阈值的趋势;建议加强关注，检查系统防水密封情况。')
        else: document.add_paragraph('预测水分未来仍处于正常范围;')         
        document.add_paragraph('温度变化趋势正常。')
    except Exception as e:
        log(e)
        word_tag=word_tag-1     
    try:       
        document.add_picture(cur_path+'//png//forcast_func_黏度.png',width=Inches(6))  
        paragraph =document.add_paragraph('图(a) 黏度预测情况') 
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER                 
    except Exception as e:
         log(e)     
    try:       
        document.add_picture(cur_path+'//png//forcast_func_温度.png',width=Inches(6))  
        paragraph =document.add_paragraph('图(b) 温度预测情况') 
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER             
    except Exception as e:  
        log(e)       
    try:       
        document.add_picture(cur_path+'//png//forcast_func_油温.png',width=Inches(6)) 
        paragraph =document.add_paragraph('图(b) 温度预测情况') 
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER            
    except Exception as e:
        log(e)                 
    try:       
        document.add_picture(cur_path+'//png//forcast_func_水活性.png',width=Inches(6)) 
        paragraph =document.add_paragraph('图(c) 水活性预测情况') 
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER   
        paragraph =document.add_paragraph('图%s 各主要属性预测情况'%word_tag)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER           
    except Exception as e:
        word_tag=word_tag-1   
        log(e)         

    document.add_heading('综合分析', level=1)  
    try:  
        word_tag=word_tag+1
        health_Score_plt(df,vis,arg_x=arg_x,arg_y=arg_y,arg_z=arg_z,facility_type=facility_type,oil_type=oil_type)
        tag,health_score,comment,item_score_set=health_Score(df.iloc[-1],vis,facility_type=facility_type,oil_type=oil_type)
        document.add_paragraph('图%s是%s的在用油的综合健康指数。由图可知，当前系统综合健康指数为%s。其中左侧散点图代表期间内各时刻检测数据及对应的分数，右侧为各项目得分。'%(word_tag,equip_name,health_score)) 
        document.add_picture(cur_path+'\png\health_index.png',width=Inches(6)) 
        paragraph =document.add_paragraph('图%s 综合健康情况'%word_tag) 
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER                
    except Exception as e:
        word_tag=word_tag-1   
        log(e)             
    document.save(equip_name+'_报告demo.docx')  
    print('报告完成')
    


    

if __name__ == '__main__':
    time1=time.time()
    #datafile_to_word('立磨主减速机.csv',data_deal=True,method='nearest',period='6H')
    project_path = os.path.split(os.path.realpath(__file__))[0]

    df_data=pd.read_csv(project_path+'\data\K301A.csv',encoding='gbk')  
    df2=df_data###建议使用100个数据行    
    #datafile_to_word2(df2,'立磨主减速机',32,data_deal=True,method='linear',period='12H',facility_type='汽轮机',oil_type='rest')
    datafile_to_word2(df2,'Claus风机131-K-301A',46,facility_type='压缩机',oil_type='rest',arg_x='黏度',arg_y='水活性',arg_z='NAS等级')
    time2=time.time()  
    print('运行时间：%ds'%(time2-time1))
    
    
    
    
    
    
    
    